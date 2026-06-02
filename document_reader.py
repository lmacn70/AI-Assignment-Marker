from pathlib import Path
from docx import Document
import fitz
from PIL import Image
import pytesseract
import io

# Tell pytesseract where Tesseract is installed
pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

def read_word_document(file_path):
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def read_pdf_document(file_path):
    text_parts = []

    with fitz.open(file_path) as pdf:
        print(f"PDF pages found: {len(pdf)}")

        for page_number, page in enumerate(pdf, start=1):
            print(f"Processing PDF page {page_number}")

            page_text = page.get_text().strip()

            if page_text:
                text_parts.append(f"\n--- Page {page_number} text ---\n")
                text_parts.append(page_text)
            else:
                text_parts.append(f"\n--- Page {page_number} OCR ---\n")

                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
                image_path = f"ocr_page_{page_number}.png"
                pix.save(image_path)

                ocr_text = pytesseract.image_to_string(
                    Image.open(image_path),
                    config="--psm 6"
                )

                if ocr_text.strip():
                    text_parts.append(ocr_text.strip())
                else:
                    text_parts.append("[No OCR text found on this page]")

    return "\n".join(text_parts)


def read_text_document(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def read_document(file_path):
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == ".docx":
        return read_word_document(file_path)

    if extension == ".pdf":
        return read_pdf_document(file_path)

    if extension == ".txt":
        return read_text_document(file_path)

    raise ValueError(f"Unsupported file type: {extension}")

