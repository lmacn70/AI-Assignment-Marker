from pathlib import Path
from docx import Document


def safe_filename(name):
    invalid = '<>:"/\\|?*'
    for char in invalid:
        name = name.replace(char, "_")
    return name


def write_feedback_document(marking_result, output_folder):
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    student_name = marking_result.get("student_name", "Student")
    safe_name = safe_filename(student_name)

    doc = Document()

    doc.add_heading("Assignment Feedback", level=1)
    doc.add_paragraph(f"Student: {student_name}")

    doc.add_heading("Criteria Results", level=2)

    criteria_results = marking_result.get("criteria_results", [])

    if criteria_results:
        for criterion in criteria_results:
            criterion_name = criterion.get("criterion", "Criterion")
            grade = criterion.get("grade", "No grade")
            feedback = criterion.get("feedback", "No feedback provided.")

            doc.add_heading(criterion_name, level=3)
            doc.add_paragraph(f"Grade: {grade}")
            doc.add_paragraph(f"Feedback: {feedback}")
    else:
        doc.add_paragraph("No criteria results were returned.")

    doc.add_heading("Overall Result", level=2)
    doc.add_paragraph(f"Overall Grade: {marking_result.get('overall_grade', 'No grade')}")
    doc.add_paragraph(f"Overall Feedback: {marking_result.get('overall_feedback', 'No feedback')}")

    output_path = output_folder / f"{safe_name}_Feedback.docx"
    doc.save(output_path)

    return output_path


def write_draft_feedback_document(feedback_result, output_folder):
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    student_name = feedback_result.get("student_name", "Student")
    safe_name = safe_filename(student_name)

    doc = Document()

    doc.add_heading("Draft Assignment Feedback", level=1)
    doc.add_paragraph(f"Student: {student_name}")
    doc.add_paragraph("This is draft feedback only. No grade has been awarded.")

    doc.add_heading("Criteria Feedback", level=2)

    criteria_feedback = feedback_result.get("criteria_feedback", [])

    if criteria_feedback:
        for item in criteria_feedback:
            criterion_name = item.get("criterion", "Criterion")
            done_well = item.get("done_well", "No strengths identified.")
            needs_to_do = item.get("needs_to_do", "No improvement advice provided.")

            doc.add_heading(criterion_name, level=3)
            doc.add_paragraph("What you have done well:")
            doc.add_paragraph(done_well)
            doc.add_paragraph("What you still need to do:")
            doc.add_paragraph(needs_to_do)
    else:
        doc.add_paragraph("No criteria feedback was returned.")

    doc.add_heading("Overall Draft Feedback", level=2)
    doc.add_paragraph(feedback_result.get("overall_feedback", "No overall feedback provided."))

    output_path = output_folder / f"{safe_name}_Draft_Feedback.docx"
    doc.save(output_path)

    return output_path

